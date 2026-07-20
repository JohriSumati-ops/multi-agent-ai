"""
document_processing/parsers/docx_parser.py

WHY THIS FILE EXISTS
---------------------
`python-docx` exposes a Word document as a sequence of `Paragraph` objects
rather than a flat string — this adapter is responsible for joining them
back into the same `raw_text` shape every other parser produces, while
preserving paragraph boundaries (`\n\n`) since Phase 2's paragraph-based
chunking strategy depends on them being recoverable.
"""

from __future__ import annotations

import docx
from docx.opc.exceptions import PackageNotFoundError

from core.exceptions import CorruptedDocumentError, EmptyDocumentError
from document_processing.parsers.base_parser import BaseParser, ParsedDocument


class DOCXParser(BaseParser):
    def parse(self, file_path: str) -> ParsedDocument:
        try:
            document = docx.Document(file_path)
        except (PackageNotFoundError, KeyError, ValueError) as exc:
            raise CorruptedDocumentError(
                "The Word document could not be opened — it may be corrupted or not a valid .docx file.",
                details={"underlying_error": str(exc)},
            ) from exc

        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        raw_text = "\n\n".join(paragraphs)

        if not raw_text.strip():
            raise EmptyDocumentError("This Word document contains no readable text.")

        core_props = document.core_properties
        return ParsedDocument(
            raw_text=raw_text,
            native_title=core_props.title or None,
            native_author=core_props.author or None,
        )
