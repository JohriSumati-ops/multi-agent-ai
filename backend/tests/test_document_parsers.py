"""
tests/test_document_parsers.py

Exercises every concrete parser against real files (a reportlab-generated
PDF, a python-docx-generated DOCX, and plain text/Markdown files) plus the
Phase 2 error-handling requirements: corrupted PDFs, empty documents, and
encrypted PDFs.
"""

from __future__ import annotations

import docx
import pytest
from pypdf import PdfWriter
from reportlab.pdfgen import canvas

from core.exceptions import CorruptedDocumentError, EmptyDocumentError, EncryptedDocumentError
from document_processing.parsers.docx_parser import DOCXParser
from document_processing.parsers.factory import get_parser
from document_processing.parsers.markdown_parser import MarkdownParser
from document_processing.parsers.pdf_parser import PDFParser
from document_processing.parsers.txt_parser import TXTParser
from models.document import DocumentFormat


def _make_pdf(path: str, text: str = "Binary search trees support O(log n) lookup.") -> None:
    c = canvas.Canvas(path)
    c.drawString(100, 750, text)
    c.save()


def _make_encrypted_pdf(path: str) -> None:
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.encrypt(user_password="secret", owner_password="ownersecret")
    with open(path, "wb") as f:
        writer.write(f)


def test_txt_parser_reads_plain_text(tmp_path) -> None:
    file_path = tmp_path / "sample.txt"
    file_path.write_text("Graphs generalize trees by allowing cycles.")
    result = TXTParser().parse(str(file_path))
    assert "Graphs generalize trees" in result.raw_text


def test_txt_parser_rejects_empty_file(tmp_path) -> None:
    file_path = tmp_path / "empty.txt"
    file_path.write_text("")
    with pytest.raises(EmptyDocumentError):
        TXTParser().parse(str(file_path))


def test_markdown_parser_extracts_h1_as_title(tmp_path) -> None:
    file_path = tmp_path / "notes.md"
    file_path.write_text("# Dynamic Programming\n\nDP solves problems by combining subproblem solutions.")
    result = MarkdownParser().parse(str(file_path))
    assert result.native_title == "Dynamic Programming"
    assert "subproblem" in result.raw_text


def test_docx_parser_extracts_paragraphs(tmp_path) -> None:
    file_path = tmp_path / "notes.docx"
    document = docx.Document()
    document.add_paragraph("Hash tables provide average O(1) lookup.")
    document.add_paragraph("Collisions are handled via chaining or open addressing.")
    document.save(str(file_path))

    result = DOCXParser().parse(str(file_path))
    assert "Hash tables" in result.raw_text
    assert "Collisions" in result.raw_text


def test_docx_parser_rejects_corrupted_file(tmp_path) -> None:
    file_path = tmp_path / "corrupted.docx"
    file_path.write_bytes(b"this is not a real docx file")
    with pytest.raises(CorruptedDocumentError):
        DOCXParser().parse(str(file_path))


def test_pdf_parser_extracts_text(tmp_path) -> None:
    file_path = tmp_path / "sample.pdf"
    _make_pdf(str(file_path))
    result = PDFParser().parse(str(file_path))
    assert "Binary search trees" in result.raw_text
    assert result.page_count == 1
    assert len(result.pages) == 1


def test_pdf_parser_rejects_corrupted_file(tmp_path) -> None:
    file_path = tmp_path / "corrupted.pdf"
    file_path.write_bytes(b"%PDF-1.4 this is not actually a valid pdf structure")
    with pytest.raises(CorruptedDocumentError):
        PDFParser().parse(str(file_path))


def test_pdf_parser_rejects_encrypted_file(tmp_path) -> None:
    file_path = tmp_path / "encrypted.pdf"
    _make_encrypted_pdf(str(file_path))
    with pytest.raises(EncryptedDocumentError):
        PDFParser().parse(str(file_path))


def test_pdf_parser_rejects_empty_text_pdf(tmp_path) -> None:
    # A blank page has no extractable text at all.
    file_path = tmp_path / "blank.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with open(file_path, "wb") as f:
        writer.write(f)

    with pytest.raises(EmptyDocumentError):
        PDFParser().parse(str(file_path))


def test_factory_dispatches_to_correct_parser() -> None:
    assert isinstance(get_parser(DocumentFormat.PDF), PDFParser)
    assert isinstance(get_parser(DocumentFormat.TXT), TXTParser)
    assert isinstance(get_parser(DocumentFormat.MARKDOWN), MarkdownParser)
    assert isinstance(get_parser(DocumentFormat.DOCX), DOCXParser)
