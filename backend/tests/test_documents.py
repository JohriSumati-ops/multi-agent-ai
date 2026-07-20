"""
tests/test_documents.py

End-to-end integration tests hitting the real HTTP routes (upload, list,
get, delete, chunks), exercising the full Upload -> Validation -> Storage
-> Parsing -> Metadata -> Cleaning -> Chunking -> Database pipeline through
`DocumentService`, with an in-memory SQLite database and a temp-directory
upload folder (see `tests/conftest.py::isolated_upload_dir`).
"""

from __future__ import annotations

import io

import docx
from reportlab.pdfgen import canvas

from core.config import settings


def _make_pdf_bytes(text: str = "Binary search trees support O(log n) average-case lookup.") -> bytes:
    buffer = io.BytesIO()
    c = canvas.Canvas(buffer)
    c.drawString(100, 750, text)
    c.save()
    return buffer.getvalue()


def _make_docx_bytes() -> bytes:
    buffer = io.BytesIO()
    document = docx.Document()
    document.add_paragraph("Hash tables provide average O(1) lookup performance.")
    document.add_paragraph("Collisions are typically handled via chaining or open addressing.")
    document.save(buffer)
    return buffer.getvalue()


def test_upload_txt_document_runs_full_pipeline(client, auth_headers) -> None:
    content = (
        b"A tree is a hierarchical data structure. It consists of nodes connected by edges. "
        b"Each node has a parent, except the root node."
    )
    response = client.post(
        "/api/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("trees_notes.txt", content, "text/plain")},
    )
    assert response.status_code == 201
    body = response.json()
    assert body["success"] is True

    document = body["data"]
    assert document["status"] == "chunked"
    assert document["file_format"] == "txt"
    assert document["word_count"] > 0
    assert document["language"] == "en"
    assert document["title"] == "Trees Notes"


def test_upload_markdown_document_uses_h1_as_title(client, auth_headers) -> None:
    content = b"# Dynamic Programming\n\nDP solves problems via subproblem reuse and memoization."
    response = client.post(
        "/api/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("untitled.md", content, "text/markdown")},
    )
    assert response.status_code == 201
    assert response.json()["data"]["title"] == "Dynamic Programming"


def test_upload_docx_document(client, auth_headers) -> None:
    response = client.post(
        "/api/v1/documents/upload",
        headers=auth_headers,
        files={
            "file": (
                "hashing_notes.docx",
                _make_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 201
    body = response.json()["data"]
    assert body["status"] == "chunked"
    assert body["file_format"] == "docx"


def test_upload_pdf_document_and_retrieve_chunks(client, auth_headers) -> None:
    response = client.post(
        "/api/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("bst_notes.pdf", _make_pdf_bytes(), "application/pdf")},
    )
    assert response.status_code == 201
    document = response.json()["data"]
    assert document["status"] == "chunked"
    assert document["page_count"] == 1

    chunks_response = client.get(f"/api/v1/documents/{document['id']}/chunks", headers=auth_headers)
    assert chunks_response.status_code == 200
    chunks = chunks_response.json()["data"]
    assert len(chunks) >= 1
    assert "Binary search trees" in chunks[0]["chunk_text"]
    assert chunks[0]["chunking_strategy"] == "paragraph"


def test_upload_rejects_unsupported_file_type(client, auth_headers) -> None:
    response = client.post(
        "/api/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("virus.exe", b"not a real document", "application/octet-stream")},
    )
    assert response.status_code == 415
    assert response.json()["error"]["code"] == "unsupported_file_type"


def test_upload_rejects_empty_file(client, auth_headers) -> None:
    response = client.post(
        "/api/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("empty.txt", b"", "text/plain")},
    )
    assert response.status_code == 422
    assert response.json()["error"]["code"] == "empty_document"


def test_upload_rejects_file_exceeding_size_limit(client, auth_headers, monkeypatch) -> None:
    monkeypatch.setattr(settings, "MAX_UPLOAD_SIZE_MB", 1)
    oversized_content = b"x" * (2 * 1024 * 1024)  # 2MB > 1MB limit
    response = client.post(
        "/api/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("huge.txt", oversized_content, "text/plain")},
    )
    assert response.status_code == 413
    assert response.json()["error"]["code"] == "file_too_large"


def test_upload_marks_document_failed_on_corrupted_pdf(client, auth_headers) -> None:
    response = client.post(
        "/api/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("corrupted.pdf", b"%PDF-1.4 not actually valid", "application/pdf")},
    )
    # The Document row creation succeeds (file is stored), but the pipeline
    # fails during parsing -> the request itself surfaces the error.
    assert response.status_code == 422
    assert response.json()["error"]["code"] == "corrupted_document"


def test_list_documents_returns_only_current_users_documents(client) -> None:
    # User A uploads a document.
    client.post(
        "/api/v1/auth/register",
        json={"email": "userA@example.com", "password": "password123"},
    )
    login_a = client.post(
        "/api/v1/auth/login", json={"email": "userA@example.com", "password": "password123"}
    )
    headers_a = {"Authorization": f"Bearer {login_a.json()['data']['access_token']}"}
    client.post(
        "/api/v1/documents/upload",
        headers=headers_a,
        files={"file": ("a_notes.txt", b"User A's private study notes about stacks and queues.", "text/plain")},
    )

    # User B should see an empty list.
    client.post(
        "/api/v1/auth/register",
        json={"email": "userB@example.com", "password": "password123"},
    )
    login_b = client.post(
        "/api/v1/auth/login", json={"email": "userB@example.com", "password": "password123"}
    )
    headers_b = {"Authorization": f"Bearer {login_b.json()['data']['access_token']}"}
    list_response = client.get("/api/v1/documents", headers=headers_b)

    assert list_response.status_code == 200
    assert list_response.json()["data"] == []


def test_get_document_returns_404_for_other_users_document(client) -> None:
    client.post("/api/v1/auth/register", json={"email": "owner@example.com", "password": "password123"})
    owner_login = client.post(
        "/api/v1/auth/login", json={"email": "owner@example.com", "password": "password123"}
    )
    owner_headers = {"Authorization": f"Bearer {owner_login.json()['data']['access_token']}"}
    upload_response = client.post(
        "/api/v1/documents/upload",
        headers=owner_headers,
        files={"file": ("private.txt", b"Some private study notes about binary trees.", "text/plain")},
    )
    document_id = upload_response.json()["data"]["id"]

    client.post("/api/v1/auth/register", json={"email": "intruder@example.com", "password": "password123"})
    intruder_login = client.post(
        "/api/v1/auth/login", json={"email": "intruder@example.com", "password": "password123"}
    )
    intruder_headers = {"Authorization": f"Bearer {intruder_login.json()['data']['access_token']}"}

    response = client.get(f"/api/v1/documents/{document_id}", headers=intruder_headers)
    assert response.status_code == 404


def test_delete_document_removes_it_and_its_chunks(client, auth_headers) -> None:
    upload_response = client.post(
        "/api/v1/documents/upload",
        headers=auth_headers,
        files={"file": ("temp_notes.txt", b"Temporary notes about linked lists and their operations.", "text/plain")},
    )
    document_id = upload_response.json()["data"]["id"]

    delete_response = client.delete(f"/api/v1/documents/{document_id}", headers=auth_headers)
    assert delete_response.status_code == 200
    assert delete_response.json()["data"]["deleted"] is True

    get_response = client.get(f"/api/v1/documents/{document_id}", headers=auth_headers)
    assert get_response.status_code == 404


def test_upload_requires_authentication(client) -> None:
    response = client.post(
        "/api/v1/documents/upload",
        files={"file": ("notes.txt", b"Some notes.", "text/plain")},
    )
    assert response.status_code == 401
